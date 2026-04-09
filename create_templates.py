"""
create_templates.py
-------------------
Gera os 7 modelos base .docx usados pelo AleDocs.
Execute UMA VEZ antes de iniciar o sistema:

    python create_templates.py

Os arquivos são salvos em ./docx_templates/
A logo é inserida no cabeçalho de todos os documentos.
"""

import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

LOGO_PATH = os.path.join('static', 'img', 'logo.png')

# ── Paleta da logo ────────────────────────────────────────────────────────────
COR_MARROM_ESCURO = RGBColor(0x2C, 0x15, 0x00)
COR_MARROM_MEDIO  = RGBColor(0x4A, 0x2C, 0x0A)
COR_DOURADO       = RGBColor(0xC8, 0x97, 0x3A)
COR_OLIVA         = RGBColor(0x4A, 0x6B, 0x15)
COR_TEXTO         = RGBColor(0x1E, 0x0F, 0x00)

TEMPLATES_DIR = 'docx_templates'
os.makedirs(TEMPLATES_DIR, exist_ok=True)


# ─── Helpers internos ────────────────────────────────────────────────────────

def v(nome):
    """Retorna placeholder do docxtpl: {{nome}}"""
    return '{{' + nome + '}}'


def margem(doc, top=3.0, bottom=2.0, left=3.0, right=2.0):
    for sec in doc.sections:
        sec.top_margin = Cm(top)
        sec.bottom_margin = Cm(bottom)
        sec.left_margin = Cm(left)
        sec.right_margin = Cm(right)


def remove_table_borders(tbl):
    tbl_pr = tbl._tbl.tblPr
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def cabecalho_com_logo(doc, subtitulo_txt=''):
    """Cabeçalho: logo | nome empresa | linha dourada."""
    tbl = doc.add_table(rows=1, cols=2)
    remove_table_borders(tbl)
    tbl.columns[0].width = Cm(3.2)
    tbl.columns[1].width = Cm(14.8)

    cel_logo = tbl.cell(0, 0)
    p_logo = cel_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(LOGO_PATH):
        p_logo.add_run().add_picture(LOGO_PATH, width=Cm(3.36))
    else:
        p_logo.add_run('[LOGO]').font.size = Pt(9)

    cel_txt = tbl.cell(0, 1)
    p_nome = cel_txt.paragraphs[0]
    p_nome.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_nome = p_nome.add_run(v('nome_emissor'))
    r_nome.bold = True
    r_nome.font.size = Pt(13)
    r_nome.font.color.rgb = COR_MARROM_ESCURO

    p_info = cel_txt.add_paragraph()
    ri = p_info.add_run(v('linha_emissor'))
    ri.font.size = Pt(8.5)
    ri.font.color.rgb = COR_MARROM_MEDIO

    if subtitulo_txt:
        p_sub = cel_txt.add_paragraph()
        rs = p_sub.add_run(subtitulo_txt)
        rs.font.size = Pt(8.5)
        rs.font.color.rgb = COR_OLIVA
        rs.italic = True

    doc.add_paragraph()
    p_linha = doc.add_paragraph()
    p_linha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl = p_linha.add_run('─' * 90)
    rl.font.size = Pt(7)
    rl.font.color.rgb = COR_DOURADO
    doc.add_paragraph()


def titulo(doc, texto, tamanho=13):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(texto.upper())
    r.bold = True
    r.font.size = Pt(tamanho)
    r.font.color.rgb = COR_MARROM_ESCURO
    doc.add_paragraph()
    return p


def subtitulo(doc, texto):
    titulo(doc, texto, tamanho=11)


def secao_titulo(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = COR_MARROM_MEDIO
    return p


def paragrafo(doc, partes, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    for txt, bold in partes:
        r = p.add_run(txt)
        r.bold = bold
        r.font.size = Pt(11)
        r.font.color.rgb = COR_TEXTO
    return p


def campo(doc, label, variavel):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f'{label}: ')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = COR_MARROM_MEDIO
    r2 = p.add_run(v(variavel))
    r2.font.size = Pt(11)
    r2.font.color.rgb = COR_TEXTO
    return p


def item_check(doc, texto, tamanho=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f'☐  {texto}')
    r.font.size = Pt(tamanho)
    r.font.color.rgb = COR_TEXTO
    return p


def assinaturas(doc, label_esq, var_esq, label_dir, nome_dir):
    doc.add_paragraph()
    doc.add_paragraph()
    tbl = doc.add_table(rows=3, cols=2)
    remove_table_borders(tbl)
    dados = [
        ('_' * 42, '_' * 42),
        (v(var_esq), nome_dir),
        (label_esq, label_dir),
    ]
    for i, (esq, dir_) in enumerate(dados):
        for j, txt in enumerate([esq, dir_]):
            p = tbl.cell(i, j).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.clear()
            r = p.add_run(txt)
            r.font.size = Pt(10.5)
            r.font.color.rgb = COR_TEXTO
            if i == 2:
                r.bold = True
                r.font.color.rgb = COR_MARROM_MEDIO


def linha_local_data(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(v('cidade') + ', ' + v('data_atual'))
    r.font.size = Pt(11)
    r.font.color.rgb = COR_TEXTO
    return p


# ═════════════════════════════════════════════════════════════════════════════
# 1. TERMO DE ENTREGA DE DOCUMENTOS
# ═════════════════════════════════════════════════════════════════════════════

def criar_termo_entrega():
    doc = Document()
    margem(doc)
    cabecalho_com_logo(doc)
    titulo(doc, 'Termo de Entrega de Documentos')
    subtitulo(doc, 'Georreferenciamento de Imóvel Rural')
    doc.add_paragraph()
    linha_local_data(doc)
    doc.add_paragraph()

    paragrafo(doc, [
        ('Eu, ', False), (v('nome_cliente'), True),
        (', CPF/CNPJ: ', False), (v('cpf_cnpj'), False),
        (', profissão: ', False), (v('profissao'), False),
        (', estado civil: ', False), (v('estado_civil'), False),
        (', residente e domiciliado(a) em ', False), (v('endereco_completo'), False),
        (', telefone: ', False), (v('telefone'), False),
        (', e-mail: ', False), (v('email'), False), ('.', False),
    ])

    doc.add_paragraph()

    paragrafo(doc, [
        ('Proprietário(a) do imóvel denominado ', False),
        (v('nome_propriedade'), True),
        (', Matrícula nº ', False), (v('matricula_imovel'), False),
        (', Comarca de ', False), (v('comarca'), False),
        (', declaro ter entregue nesta data os seguintes documentos para a '
         'realização do serviço de Georreferenciamento:', False),
    ])

    doc.add_paragraph()
    secao_titulo(doc, 'Documentos Entregues:')

    itens = [
        'Matrícula do imóvel atualizada (emitida há no máximo 30 dias)',
        f'CCIR – Certificado de Cadastro de Imóvel Rural nº: {v("numero_ccir")}',
        'Comprovante de ITR dos últimos 5 exercícios ou Declaração de Isenção',
        f'Cópia do CPF, RG nº {v("rg")} e CNH do(a) proprietário(a)',
        f'CAR – Cadastro Ambiental Rural nº: {v("numero_car")}',
        'Escritura pública ou instrumento particular registrado',
        'Planta ou croqui anterior do imóvel (se disponível)',
        'Procuração com poderes específicos (se representado por terceiro)',
        'ART/TRT de georreferenciamento anterior (se houver)',
        'Comprovante de DARF (se aplicável)',
        f'Cópias de CPF e RG do(a) cônjuge: {v("nome_conjuge")}',
        'Outros: ___________________________________________________',
    ]
    for item in itens:
        item_check(doc, item)

    doc.add_paragraph()

    paragrafo(doc, [
        ('Declaro que os documentos acima listados foram entregues voluntariamente para a realização '
         'do serviço contratado. Os originais serão devolvidos ao término do serviço.', False),
    ])

    doc.add_paragraph()
    secao_titulo(doc, 'Observações:')
    p_obs = doc.add_paragraph()
    p_obs.add_run(v('observacoes')).font.size = Pt(12)

    assinaturas(
        doc,
        'CONTRATANTE / PROPRIETÁRIO(A)', 'nome_cliente',
        'RESPONSÁVEL TÉCNICO / ADVOGADO(A)', v('nome_emissor'),
    )

    path = os.path.join(TEMPLATES_DIR, 'termo_entrega.docx')
    doc.save(path)
    print(f'  ✓  {path}')


# ═════════════════════════════════════════════════════════════════════════════
# 2. CHECKLIST DE AVERBAÇÃO DE GEORREFERENCIAMENTO
# ═════════════════════════════════════════════════════════════════════════════

def criar_checklist_averbacao():
    doc = Document()
    margem(doc)
    cabecalho_com_logo(doc)
    titulo(doc, 'Checklist de Averbação de Georreferenciamento')
    subtitulo(doc, 'Imóvel Rural – Lei nº 10.267/2001 | Art. 176 §3º da Lei 6.015/73')
    doc.add_paragraph()

    for lbl, var_nome in [
        ('Cliente', 'nome_cliente'), ('CPF/CNPJ', 'cpf_cnpj'),
        ('RG / Identidade', 'rg'),
        ('Cônjuge', 'nome_conjuge'),
        ('Imóvel', 'nome_propriedade'), ('Matrícula', 'matricula_imovel'),
        ('Cód. INCRA / SNCR', 'codigo_incra'),
        ('Nº CCIR', 'numero_ccir'),
        ('Nº CAR', 'numero_car'),
        ('Comarca', 'comarca'), ('Área Georreferenciada', 'area_imovel'),
        ('Área Registrada', 'area_registrada'),
        ('Perímetro', 'perimetro_imovel'),
        ('Protocolo SIGEF', 'numero_sigef'),
        ('Data', 'data_atual'),
    ]:
        campo(doc, lbl, var_nome)

    doc.add_paragraph()

    secoes = [
        ('1. Documentação do Imóvel', [
            'Matrícula atualizada (emitida há no máximo 30 dias)',
            'Certidão de inteiro teor da matrícula',
            'CCIR – Certificado de Cadastro de Imóvel Rural (atual)',
            'CAR – Cadastro Ambiental Rural (ativo)',
            'ITR dos últimos 5 exercícios ou Declaração de Isenção',
            'Escritura pública ou instrumento particular registrado',
        ]),
        ('2. Documentação do Proprietário', [
            'CPF e documento de identidade (RG/CNH)',
            'Comprovante de residência atualizado',
            'Certidão de estado civil (casamento/nascimento)',
            'Procuração (se representado por terceiro)',
            'CNPJ e contrato social (se pessoa jurídica)',
        ]),
        ('3. Documentação Técnica – Georreferenciamento', [
            'Planta com coordenadas georreferenciadas (SIRGAS 2000)',
            'Memorial descritivo assinado pelo responsável técnico',
            'ART/TRT do responsável técnico (CREA)',
            'Relatório técnico de georreferenciamento',
            'Arquivo digital (KML/Shapefile/XML) para envio ao SIGEF',
            'Número de protocolo SIGEF (após aprovação)',
            'Certificação pelo INCRA (SIGEF/SNCR)',
        ]),
        ('4. Procedimentos no Cartório de Registro de Imóveis', [
            'Protocolo dos documentos no Cartório',
            'Guia de recolhimento de emolumentos paga',
            'Notificação de confrontantes realizada',
            'Prazo de impugnação aguardado (15 dias – Lei nº 10.267/2001)',
            'Manifestação de anuência dos confrontantes anexada',
            'Averbação efetivada na matrícula do imóvel',
        ]),
        ('5. Providências Pós-Averbação', [
            'Cópia da matrícula com a averbação do georreferenciamento',
            'Atualização do CCIR junto ao INCRA',
            'Atualização do CAR (se necessário)',
            'Entrega dos documentos finais ao cliente',
        ]),
    ]

    for sec_titulo, itens in secoes:
        secao_titulo(doc, sec_titulo)
        for item in itens:
            item_check(doc, item)
        doc.add_paragraph()

    secao_titulo(doc, 'Observações:')
    p_obs = doc.add_paragraph()
    p_obs.add_run(v('observacoes')).font.size = Pt(12)
    for _ in range(3):
        p = doc.add_paragraph()
        p.add_run('_' * 95).font.size = Pt(9)

    doc.add_paragraph()
    p_resp = doc.add_paragraph()
    p_resp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_resp.add_run('Responsável: ________________________  Data: ' + v('data_curta')).font.size = Pt(11)

    path = os.path.join(TEMPLATES_DIR, 'checklist_averbacao.docx')
    doc.save(path)
    print(f'  ✓  {path}')


# ═════════════════════════════════════════════════════════════════════════════
# 3. CONTRATO DE GEORREFERENCIAMENTO
# ═════════════════════════════════════════════════════════════════════════════

def criar_contrato_georreferenciamento():
    doc = Document()
    margem(doc)
    cabecalho_com_logo(doc)
    titulo(doc, 'Contrato de Prestação de Serviços de Georreferenciamento')
    doc.add_paragraph()
    linha_local_data(doc)
    doc.add_paragraph()

    # Partes
    secao_titulo(doc, 'Das Partes')

    paragrafo(doc, [
        ('CONTRATANTE: ', True), (v('nome_cliente'), False),
        (', ', False), (v('estado_civil'), False),
        (', ', False), (v('profissao'), False),
        (', CPF/CNPJ nº ', False), (v('cpf_cnpj'), False),
        (', residente e domiciliado(a) em ', False),
        (v('endereco_completo'), False), (', telefone: ', False),
        (v('telefone'), False), (', e-mail: ', False), (v('email'), False), ('.', False),
    ])

    doc.add_paragraph()

    paragrafo(doc, [
        ('CONTRATADA: ', True),
        (v('nome_emissor'), True),
        (', profissão: ', False), (v('profissao_emissor'), False),
        (', Registro: ', False), (v('registro_emissor'), False),
        (', endereço: ', False), (v('endereco_emissor'), False),
        (', telefone: ', False), (v('telefone_emissor'), False),
        (', e-mail: ', False), (v('email_emissor'), False), ('.', False),
    ])

    doc.add_paragraph()
    secao_titulo(doc, 'Cláusula 1ª – Do Objeto')
    paragrafo(doc, [
        ('O presente contrato tem por objeto a prestação de serviços técnicos de '
         'georreferenciamento do imóvel rural denominado ', False),
        (v('nome_propriedade'), True),
        (', Matrícula nº ', False), (v('matricula_imovel'), False),
        (', Código INCRA/SNCR: ', False), (v('codigo_incra'), False),
        (', CCIR nº ', False), (v('numero_ccir'), False),
        (', CAR nº ', False), (v('numero_car'), False),
        (', com área registrada de ', False), (v('area_registrada'), False),
        (', área georreferenciada de ', False), (v('area_imovel'), False),
        (', perímetro de ', False), (v('perimetro_imovel'), False),
        (', localizado no município de ', False), (v('municipio_uf_imovel'), False),
        (', Comarca de ', False), (v('comarca'), False),
        (', cujas confrontações são: Norte/Nordeste – ', False), (v('confrontante_norte'), False),
        ('; Sul/Sudoeste – ', False), (v('confrontante_sul'), False),
        ('; Leste – ', False), (v('confrontante_leste'), False),
        ('; Oeste – ', False), (v('confrontante_oeste'), False),
        ('; nos termos da Lei Federal nº 10.267/2001 e do Decreto nº 4.449/2002, '
         'com vistas à certificação pelo INCRA e posterior averbação no Cartório de Registro de Imóveis.', False),
    ])

    doc.add_paragraph()
    secao_titulo(doc, 'Cláusula 2ª – Das Obrigações da Contratada')
    obrigacoes_contratada = [
        'Executar os serviços de campo (levantamento topográfico georreferenciado) com precisão exigida pelo INCRA;',
        'Elaborar o Memorial Descritivo e a Planta do Imóvel conforme normas técnicas vigentes;',
        'Registrar a ART/TRT junto ao órgão de Registro competente, responsabilizando-se tecnicamente pelo serviço;',
        'Protocolar os arquivos digitais no SIGEF/INCRA e acompanhar o processo de certificação;',
        'Fornecer ao Contratante todos os documentos gerados após a conclusão do serviço.',
    ]
    for ob in obrigacoes_contratada:
        item_check(doc, ob)

    doc.add_paragraph()
    secao_titulo(doc, 'Cláusula 3ª – Das Obrigações do Contratante')
    obrigacoes_contratante = [
        'Fornecer toda a documentação solicitada no prazo acordado;',
        'Permitir o acesso ao imóvel para realização dos serviços de campo;',
        'Comunicar previamente os confrontantes sobre os trabalhos de campo;',
        'Efetuar o pagamento nas condições estipuladas na Cláusula 4ª.',
    ]
    for ob in obrigacoes_contratante:
        item_check(doc, ob)

    doc.add_paragraph()
    secao_titulo(doc, 'Cláusula 4ª – Do Preço e Forma de Pagamento')
    paragrafo(doc, [
        ('Pelos serviços ora contratados, o CONTRATANTE pagará à CONTRATADA o valor de '
         'R$ ______________ (______________________________), nas seguintes condições: '
         '_________________________________________________________________________.', False),
    ])

    doc.add_paragraph()
    secao_titulo(doc, 'Cláusula 5ª – Do Prazo')
    paragrafo(doc, [
        ('Os serviços serão executados no prazo de ______ (______) dias corridos, contados '
         'da data de assinatura deste instrumento e do recebimento integral da documentação '
         'listada no Termo de Entrega de Documentos.', False),
    ])

    doc.add_paragraph()
    secao_titulo(doc, 'Cláusula 6ª – Da Rescisão')
    paragrafo(doc, [
        ('O presente contrato poderá ser rescindido por qualquer das partes, mediante '
         'comunicação prévia de 15 (quinze) dias, ficando a parte que der causa à rescisão '
         'responsável pelos prejuízos comprovadamente causados.', False),
    ])

    doc.add_paragraph()
    secao_titulo(doc, 'Cláusula 7ª – Do Foro')
    paragrafo(doc, [
        ('As partes elegem o foro da Comarca de ', False), (v('comarca'), False),
        (' para dirimir quaisquer litígios oriundos do presente instrumento, '
         'com renúncia expressa a qualquer outro, por mais privilegiado que seja.', False),
    ])

    doc.add_paragraph()
    paragrafo(doc, [
        ('E por estarem assim justos e contratados, assinam o presente instrumento '
         'em 02 (duas) vias de igual teor e forma, na presença de 02 (duas) testemunhas.', False),
    ])

    doc.add_paragraph()
    linha_local_data(doc)

    assinaturas(
        doc,
        'CONTRATANTE', 'nome_cliente',
        'CONTRATADA / RESPONSÁVEL TÉCNICO', v('nome_emissor'),
    )

    doc.add_paragraph()
    secao_titulo(doc, 'Testemunhas:')
    tbl = doc.add_table(rows=3, cols=2)
    for i, (txt_e, txt_d) in enumerate([
        ('_' * 40, '_' * 40),
        ('Nome: ___________________________', 'Nome: ___________________________'),
        ('CPF: ____________________________', 'CPF: ____________________________'),
    ]):
        for j, txt in enumerate([txt_e, txt_d]):
            p = tbl.cell(i, j).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(txt).font.size = Pt(11)

    path = os.path.join(TEMPLATES_DIR, 'contrato_georreferenciamento.docx')
    doc.save(path)
    print(f'  ✓  {path}')


# ═════════════════════════════════════════════════════════════════════════════
# 4. MEMORIAL DESCRITIVO
# ═════════════════════════════════════════════════════════════════════════════

def criar_memorial_descritivo():
    doc = Document()
    margem(doc)
    cabecalho_com_logo(doc)
    titulo(doc, 'Memorial Descritivo')
    subtitulo(doc, 'Levantamento Topográfico Georreferenciado – Imóvel Rural')
    doc.add_paragraph()
    linha_local_data(doc)
    doc.add_paragraph()

    secao_titulo(doc, '1. Identificação do Imóvel')
    for lbl, var_nome in [
        ('Denominação', 'nome_propriedade'),
        ('Matrícula nº (CRI)', 'matricula_imovel'),
        ('Código INCRA / SNCR', 'codigo_incra'),
        ('Nº do CCIR', 'numero_ccir'),
        ('Nº do CAR', 'numero_car'),
        ('Comarca', 'comarca'),
        ('Município / UF', 'municipio_uf_imovel'),
        ('Área Registrada (Escritura)', 'area_registrada'),
        ('Área Georreferenciada', 'area_imovel'),
        ('Perímetro', 'perimetro_imovel'),
        ('Sistema de Referência', 'fuso_sirgas'),
        ('Protocolo SIGEF', 'numero_sigef'),
    ]:
        campo(doc, lbl, var_nome)

    doc.add_paragraph()
    secao_titulo(doc, '2. Identificação do Proprietário')
    for lbl, var_nome in [
        ('Nome', 'nome_cliente'),
        ('CPF/CNPJ', 'cpf_cnpj'),
        ('RG / Identidade', 'rg'),
        ('Profissão', 'profissao'),
        ('Estado Civil', 'estado_civil'),
        ('Cônjuge', 'nome_conjuge'),
        ('CPF do Cônjuge', 'cpf_conjuge'),
        ('Endereço', 'endereco_completo'),
        ('Telefone', 'telefone'),
        ('E-mail', 'email'),
    ]:
        campo(doc, lbl, var_nome)

    doc.add_paragraph()
    secao_titulo(doc, '3. Confrontações do Imóvel')
    for lbl, var_nome in [
        ('Confrontante ao Norte / Nordeste', 'confrontante_norte'),
        ('Confrontante ao Sul / Sudoeste', 'confrontante_sul'),
        ('Confrontante a Leste / Este', 'confrontante_leste'),
        ('Confrontante a Oeste / Poente', 'confrontante_oeste'),
    ]:
        campo(doc, lbl, var_nome)

    doc.add_paragraph()
    secao_titulo(doc, '4. Descrição dos Limites (Planilha de Vértices)')
    paragrafo(doc, [
        ('Inicia-se a descrição deste imóvel no vértice M-01, de coordenadas '
         'N=_____________ e E=_____________ (', False),
        (v('fuso_sirgas'), False),
        ('), localizado na divisa com ', False),
        (v('confrontante_norte'), False),
        ('; deste, segue confrontando com os seguintes limites:', False),
    ])
    doc.add_paragraph()

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    cabecalhos = ['Vértice', 'Azimute', 'Distância (m)', 'Confrontação', 'Coordenadas E / N']
    for i, txt in enumerate(cabecalhos):
        p = tbl.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(10)
    for _ in range(8):
        row = tbl.add_row()
        for cel in row.cells:
            cel.paragraphs[0].add_run('').font.size = Pt(10)

    doc.add_paragraph()
    secao_titulo(doc, '5. Dados Quantitativos')
    paragrafo(doc, [
        ('Área georreferenciada: ', True), (v('area_imovel'), False), ('\n', False),
        ('Área registrada (matrícula nº ', True), (v('matricula_imovel'), False), ('): ', True),
        (v('area_registrada'), False), ('\n', False),
        ('Perímetro: ', True), (v('perimetro_imovel'), False),
    ])

    doc.add_paragraph()
    secao_titulo(doc, '6. Equipamento e Metodologia Utilizados')
    paragrafo(doc, [
        ('Equipamento: ____________________________________________\n'
         'Método de posicionamento: _______________________________\n'
         'Software de processamento: ______________________________\n'
         'Sistema de referência: ', False),
        (v('fuso_sirgas'), False),
    ])

    doc.add_paragraph()
    secao_titulo(doc, '6. Declaração do Responsável Técnico')
    paragrafo(doc, [
        ('Declaro que o presente Memorial Descritivo foi elaborado com base em levantamento '
         'topográfico georreferenciado executado in loco, estando de acordo com as normas do '
         'INCRA e da ABNT NBR 14.166, responsabilizando-me técnica e legalmente pelas '
         'informações aqui prestadas.', False),
    ])

    doc.add_paragraph()
    linha_local_data(doc)
    doc.add_paragraph()

    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sig.add_run('_' * 50).font.size = Pt(12)
    p_nome = doc.add_paragraph()
    p_nome.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nome.add_run(v('nome_emissor') + '\n' + v('registro_emissor')).font.size = Pt(11)

    path = os.path.join(TEMPLATES_DIR, 'memorial_descritivo.docx')
    doc.save(path)
    print(f'  ✓  {path}')


# ═════════════════════════════════════════════════════════════════════════════
# 5. RECIBO DE PRESTAÇÃO DE SERVIÇO
# ═════════════════════════════════════════════════════════════════════════════

def criar_recibo_servico():
    doc = Document()
    margem(doc)
    cabecalho_com_logo(doc)
    titulo(doc, 'Recibo de Prestação de Serviço')
    doc.add_paragraph()
    linha_local_data(doc)
    doc.add_paragraph()

    paragrafo(doc, [
        ('Recebo(emos) de ', False), (v('nome_cliente'), True),
        (', CPF/CNPJ nº ', False), (v('cpf_cnpj'), False),
        (', RG nº ', False), (v('rg'), False),
        (', residente em ', False), (v('endereco_completo'), False),
        (', a importância de ', False),
        (v('valor_servico'), True),
        (' referente à prestação dos seguintes serviços:', False),
    ])

    doc.add_paragraph()
    secao_titulo(doc, 'Forma de Pagamento')
    paragrafo(doc, [
        (v('forma_pagamento'), False),
    ])

    doc.add_paragraph()
    paragrafo(doc, [
        ('Serviços referentes ao imóvel rural denominado ', False),
        (v('nome_propriedade'), True),
        (', Matrícula nº ', False), (v('matricula_imovel'), False),
        (', Comarca de ', False), (v('comarca'), False), (':', False),
    ])

    doc.add_paragraph()

    for item in [
        'Serviços técnicos de georreferenciamento',
        'Elaboração de Memorial Descritivo e Planta',
        'Protocolo e acompanhamento junto ao INCRA/SIGEF',
        'Averbação no Cartório de Registro de Imóveis',
        'Assessoria jurídica e documental',
        'Outros: _______________________________________________',
    ]:
        item_check(doc, item)

    doc.add_paragraph()
    paragrafo(doc, [
        ('Forma de Pagamento: ', True),
        ('(  ) Dinheiro  (  ) PIX  (  ) Transferência  (  ) Cheque  (  ) Outro: ________', False),
    ])

    doc.add_paragraph()
    paragrafo(doc, [
        ('Para maior clareza e validade, firmo(amos) o presente recibo, dando plena quitação '
         'do valor acima descrito.', False),
    ])

    doc.add_paragraph()
    linha_local_data(doc)
    doc.add_paragraph()

    assinaturas(
        doc,
        'PRESTADOR DE SERVIÇO', 'nome_emissor',
        'CLIENTE / CONTRATANTE', v('nome_cliente'),
    )

    path = os.path.join(TEMPLATES_DIR, 'recibo_servico.docx')
    doc.save(path)
    print(f'  ✓  {path}')


# ═════════════════════════════════════════════════════════════════════════════
# 6. CONTRATO DE SERVIÇOS ADVOCATÍCIOS E TÉCNICOS – DAÇÃO EM PAGAMENTO
# ═════════════════════════════════════════════════════════════════════════════

def criar_contrato_dacao():
    doc = Document()
    margem(doc)
    cabecalho_com_logo(doc)
    titulo(doc, 'Contrato de Prestação de Serviços Advocatícios e Técnicos')
    subtitulo(doc, 'Com Cláusula de Dação em Pagamento')
    doc.add_paragraph()
    linha_local_data(doc)
    doc.add_paragraph()

    secao_titulo(doc, 'Das Partes')

    paragrafo(doc, [
        ('CONTRATANTE: ', True), (v('nome_cliente'), False),
        (', ', False), (v('estado_civil'), False),
        (', ', False), (v('profissao'), False),
        (', CPF/CNPJ nº ', False), (v('cpf_cnpj'), False),
        (', residente e domiciliado(a) em ', False),
        (v('endereco_completo'), False),
        (', telefone: ', False), (v('telefone'), False),
        (', e-mail: ', False), (v('email'), False), ('.', False),
    ])

    doc.add_paragraph()

    paragrafo(doc, [
        ('CONTRATADA: ', True),
        (v('nome_emissor'), True),
        (', profissão: ', False), (v('profissao_emissor'), False),
        (', Registro: ', False), (v('registro_emissor'), False),
        (', endereço: ', False), (v('endereco_emissor'), False),
        (', telefone: ', False), (v('telefone_emissor'), False),
        (', e-mail: ', False), (v('email_emissor'), False), ('.', False),
    ])

    doc.add_paragraph()
    secao_titulo(doc, 'Cláusula 1ª – Do Objeto')
    paragrafo(doc, [
        ('O presente instrumento tem por objeto a prestação conjunta de serviços advocatícios '
         '(regularização fundiária, averbação de georreferenciamento e demais atos registrais) '
         'e serviços técnicos de georreferenciamento relativos ao imóvel rural denominado ', False),
        (v('nome_propriedade'), True),
        (', Matrícula nº ', False), (v('matricula_imovel'), False),
        (', com área de ', False), (v('area_imovel'), False),
        (', Comarca de ', False), (v('comarca'), False), ('.', False),
    ])

    doc.add_paragraph()
    secao_titulo(doc, 'Cláusula 2ª – Da Remuneração e da Dação em Pagamento')
    paragrafo(doc, [
        ('Em contrapartida aos serviços ora contratados, cujo valor total ficou estimado em '
         'R$ ______________ (______________________________), as partes, de comum acordo, '
         'ajustam que o pagamento será realizado mediante DAÇÃO EM PAGAMENTO, consistente na '
         'transferência à CONTRATADA de FRAÇÃO IDEAL de ______ ha (________ hectares), '
         'correspondente a ______% da área total do imóvel acima referido, '
         'após a conclusão e averbação do georreferenciamento.', False),
    ])

    doc.add_paragraph()
    paragrafo(doc, [
        ('§ 1º – A fração do imóvel a ser transferida está localizada: '
         '_______________________________________________________________', False),
    ])
    doc.add_paragraph()
    paragrafo(doc, [
        ('§ 2º – As partes reconhecem que o valor da fração entregue é equivalente e suficiente '
         'para quitar integralmente as obrigações decorrentes do presente contrato, '
         'não havendo qualquer diferença a receber ou restituir.', False),
    ])
    doc.add_paragraph()
    paragrafo(doc, [
        ('§ 3º – As despesas com a lavratura e registro da escritura de dação em pagamento, '
         'emolumentos cartorários, ITBI e demais encargos ficarão a cargo de: '
         '(  ) CONTRATANTE  (  ) CONTRATADA  (  ) Rateio entre as partes.', False),
    ])

    doc.add_paragraph()
    secao_titulo(doc, 'Cláusula 3ª – Das Obrigações das Partes')
    paragrafo(doc, [
        ('I – CONTRATADA obriga-se a: executar os serviços com diligência e dentro dos prazos '
         'legais; manter sigilo sobre as informações do cliente; apresentar relatórios periódicos '
         'sobre o andamento dos serviços; arcar com as custas processuais nas ações judiciais '
         'eventualmente propostas.\n\n'
         'II – CONTRATANTE obriga-se a: fornecer toda a documentação necessária; '
         'facilitar o acesso ao imóvel; não realizar atos que possam prejudicar o andamento '
         'dos serviços; colaborar com as notificações de confrontantes.', False),
    ])

    doc.add_paragraph()
    secao_titulo(doc, 'Cláusula 4ª – Da Vigência')
    paragrafo(doc, [
        ('O presente contrato vigorará até o cumprimento integral de seu objeto, '
         'incluindo a conclusão da transferência da fração imóvel acordada, '
         'podendo ser prorrogado mediante aditivo assinado pelas partes.', False),
    ])

    doc.add_paragraph()
    secao_titulo(doc, 'Cláusula 5ª – Da Rescisão')
    paragrafo(doc, [
        ('Em caso de rescisão por culpa do CONTRATANTE, serão devidos honorários '
            'proporcionais ao trabalho realizado, apurados conforme a Tabela de Registro Profissional aplicável. '
         'Em caso de rescisão por culpa da CONTRATADA, será restituída ao CONTRATANTE '
         'a documentação entregue, sem ônus.', False),
    ])

    doc.add_paragraph()
    secao_titulo(doc, 'Cláusula 6ª – Do Foro')
    paragrafo(doc, [
        ('Fica eleito o foro da Comarca de ', False), (v('comarca'), False),
        (' para dirimir quaisquer divergências oriundas deste instrumento.', False),
    ])

    doc.add_paragraph()
    paragrafo(doc, [
        ('Assim, por estarem de pleno acordo, firmam o presente instrumento particular '
         'em 02 (duas) vias de igual teor e forma, juntamente com 02 (duas) testemunhas.', False),
    ])

    doc.add_paragraph()
    linha_local_data(doc)

    assinaturas(
        doc,
        'CONTRATANTE', 'nome_cliente',
        'CONTRATADA / RESPONSÁVEL', v('nome_emissor'),
    )

    doc.add_paragraph()
    secao_titulo(doc, 'Testemunhas:')
    for _ in range(2):
        p = doc.add_paragraph()
        p.add_run('Nome: _______________________________________  CPF: ___________________').font.size = Pt(11)

    path = os.path.join(TEMPLATES_DIR, 'contrato_dacao.docx')
    doc.save(path)
    print(f'  ✓  {path}')


# ═════════════════════════════════════════════════════════════════════════════
# 7. PROCURAÇÃO PARTICULAR
# ═════════════════════════════════════════════════════════════════════════════

def criar_procuracao():
    doc = Document()
    margem(doc)
    cabecalho_com_logo(doc, 'Documento Jurídico – Outorga de Poderes')
    titulo(doc, 'Procuração Particular')
    doc.add_paragraph()
    linha_local_data(doc)
    doc.add_paragraph()

    secao_titulo(doc, 'OUTORGANTE (MANDANTE)')
    paragrafo(doc, [
        (v('nome_cliente'), True),
        (', ', False), (v('estado_civil'), False), (', ', False), (v('profissao'), False),
        (', portador(a) do CPF nº ', False), (v('cpf_cnpj'), False),
        (', residente e domiciliado(a) em ', False), (v('endereco_completo'), False),
        (', telefone: ', False), (v('telefone'), False),
        (', e-mail: ', False), (v('email'), False), (',', False),
    ])

    doc.add_paragraph()
    secao_titulo(doc, 'OUTORGADO(A) (MANDATÁRIO / PROCURADOR)')
    paragrafo(doc, [
        (v('nome_emissor'), True),
        (', ', False), (v('nacionalidade_emissor'), False),
        (', ', False), (v('estado_civil_emissor'), False),
        (', profissão: ', False), (v('profissao_emissor'), False),
        (', CPF/CNPJ nº ', False), (v('documento_emissor'), False),
        (', Registro: ', False), (v('registro_emissor'), False),
        (', residente em ', False), (v('endereco_emissor'), False), ('.', False),
    ])

    doc.add_paragraph()
    secao_titulo(doc, 'OBJETO E PODERES OUTORGADOS')
    paragrafo(doc, [
        ('Pelo presente instrumento, o(a) OUTORGANTE nomeia e constitui o(a) '
         'OUTORGADO(A) como seu(sua) bastante procurador(a), outorgando-lhe poderes '
         'específicos para, em seu nome, praticar os seguintes atos relacionados ao '
         'imóvel rural denominado ', False),
        (v('nome_propriedade'), True),
        (', Matrícula nº ', False), (v('matricula_imovel'), False),
        (', Comarca de ', False), (v('comarca'), False), (':', False),
    ])

    poderes = [
        'Representar o(a) Outorgante perante o INCRA, SIGEF, Cartório de Registro de '
        'Imóveis, Prefeituras Municipais, Secretarias de Estado e quaisquer órgãos '
        'públicos federais, estaduais ou municipais;',
        'Assinar requerimentos, declarações, formulários, atas e demais documentos '
        'necessários ao georreferenciamento e à averbação do imóvel;',
        'Protocolar processos administrativos e judiciais, acompanhar seu andamento '
        'e praticar todos os atos necessários para sua conclusão;',
        'Requerer e retirar certidões, alvarás, licenças e autorizações junto '
        'aos órgãos competentes;',
        'Assinar o processo de certificação no SIGEF/INCRA e todos os documentos '
        'técnicos relacionados;',
        'Notificar confrontantes, receber e assinar documentos de reconhecimento '
        'de limites e divisas;',
        'Negociar e firmar acordos relativos a pendências registrais ou '
        'retificações de área, em termos que julgar convenientes;',
        'Substabelecer os poderes ora outorgados, com ou sem reserva das mesmas '
        'faculdades, a critério do(a) Outorgado(a).',
    ]
    for p_txt in poderes:
        item_check(doc, p_txt)

    doc.add_paragraph()
    secao_titulo(doc, 'VALIDADE')
    paragrafo(doc, [
        ('A presente procuração tem validade de ', False),
        ('1 (um) ano', True),
        (' a contar da data de sua assinatura, podendo ser revogada a qualquer tempo '
         'mediante comunicação escrita ao(à) Outorgado(a).', False),
    ])

    doc.add_paragraph()
    secao_titulo(doc, 'DECLARAÇÃO FINAL')
    paragrafo(doc, [
        ('O(A) OUTORGANTE declara que os dados acima são verídicos e que a presente '
         'procuração é outorgada de forma livre e espontânea, para os fins expressos '
         'neste instrumento.', False),
    ])

    doc.add_paragraph()
    paragrafo(doc, [
        ('Firmam o presente instrumento em 02 (duas) vias de igual teor e forma.', False),
    ])
    doc.add_paragraph()
    linha_local_data(doc)

    assinaturas(
        doc,
        'OUTORGANTE', 'nome_cliente',
        'OUTORGADO(A) / PROCURADOR(A)', v('nome_emissor'),
    )

    doc.add_paragraph()
    secao_titulo(doc, 'TESTEMUNHAS')
    tbl_t = doc.add_table(rows=3, cols=2)
    remove_table_borders(tbl_t)
    t_dados = [
        ('_' * 42, '_' * 42),
        ('Nome: ______________________________', 'Nome: ______________________________'),
        ('CPF:  ______________________________', 'CPF:  ______________________________'),
    ]
    for ri, (esq, dir_) in enumerate(t_dados):
        for j, txt in enumerate([esq, dir_]):
            p = tbl_t.cell(ri, j).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.clear()
            r = p.add_run(txt)
            r.font.size = Pt(10.5)
            r.font.color.rgb = COR_TEXTO

    doc.add_paragraph()
    p_nota = doc.add_paragraph()
    p_nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_nota = p_nota.add_run('Reconhecimento de firma em cartório recomendado para maior segurança jurídica.')
    r_nota.font.size = Pt(9)
    r_nota.italic = True
    r_nota.font.color.rgb = COR_MARROM_MEDIO

    path = os.path.join(TEMPLATES_DIR, 'procuracao.docx')
    doc.save(path)
    print(f'  ✓  {path}')


# ═════════════════════════════════════════════════════════════════════════════
# EXECUÇÃO
# ═════════════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    print('Gerando 7 modelos .docx com logo e paleta da marca...')
    print()
    criar_termo_entrega()
    criar_checklist_averbacao()
    criar_contrato_georreferenciamento()
    criar_memorial_descritivo()
    criar_recibo_servico()
    criar_contrato_dacao()
    criar_procuracao()
    print()
    print('Pronto! 7 modelos criados em ./docx_templates/')
    print('Agora execute: python app.py')
